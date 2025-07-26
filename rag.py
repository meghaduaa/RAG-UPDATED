import os
from docx import Document as DocxDocument
import PyPDF2
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain_core.documents import Document as LCDocument
from langchain_community.llms import LlamaCpp
from langchain_core.prompts import PromptTemplate

DB_PATH = "faiss_db"
DOCS_FOLDER = "document"
MODEL_PATH = "phi-2.Q4_K_M.gguf"

def load_documents(folder_path):
    all_docs = []

    for filename in os.listdir(folder_path):
        filepath = os.path.join(folder_path, filename)
        if filename.endswith(".docx"):
            doc = DocxDocument(filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
        elif filename.endswith(".pdf"):
            text = ""
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() or ""
        elif filename.endswith(".txt"):
            with open(filepath, "r", encoding="utf-8") as f:
                text = f.read()
        else:
            continue  # Skip unsupported files

        if text.strip():
            all_docs.append(LCDocument(page_content=text))

    return all_docs

def get_text_chunks(docs):
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    return splitter.split_documents(docs)

def get_embeddings():
    return HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

def get_vectorstore(docs, embeddings):
    if os.path.exists(DB_PATH):
        return FAISS.load_local(DB_PATH, embeddings, allow_dangerous_deserialization=True)
    else:
        db = FAISS.from_documents(docs, embeddings)
        db.save_local(DB_PATH)
        return db

def get_llm():
    return LlamaCpp(
        model_path=MODEL_PATH,
        temperature=0.5,
        max_tokens=1024,
        top_p=1,
        n_ctx=2048,
        n_batch=512,
        verbose=False
    )

def build_prompt():
    return PromptTemplate.from_template("""
Use the following context to answer the question. 
If you don't know the answer, just say "I don't know".

Context:
{context}

Question:
{question}
""")

def main():
    print("Loading and processing documents...")
    docs = load_documents(DOCS_FOLDER)
    chunks = get_text_chunks(docs)

    print("Embedding and storing in FAISS...")
    embeddings = get_embeddings()
    vectordb = get_vectorstore(chunks, embeddings)

    print("Loading Phi-2 model...")
    llm = get_llm()

    print("Creating QA chain...")
    qa = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=vectordb.as_retriever(),
        chain_type="stuff",
        chain_type_kwargs={"prompt": build_prompt()}
    )

    print("\nAsk a question (type 'exit' to quit)")
    while True:
        query = input(">> ")
        if query.lower() == "exit":
            break
        answer = qa.invoke({"query": query})["result"]
        print("\nAnswer:", answer, "\n")

if __name__ == "__main__":
    main()
